import telebot
from telebot.types import ReplyKeyboardMarkup, KeyboardButton
from datetime import datetime
from docx import Document

API_TOKEN = '7158427390:AAGmiMKa3Fx2TW01-ANC_KHTPSZ9VTc3KDA'

bot = telebot.TeleBot(API_TOKEN)

# 14 ta lift nomlari
lifts = [
    "Lift 1", "Lift 2", "Lift 3", "Lift 4", "Lift 5", 
    "Lift 6", "Lift 7", "Lift 8", "Lift 9", "Lift 10", 
    "Lift 11", "Lift 12", "Lift 13", "Lift 14"
]

# Navbatchilar soni
shifts = ["Navbatchi 1", "Navbatchi 2", "Navbatchi 3"]

# Har bir navbatchi kiritgan ma'lumotlarni saqlash uchun
lift_statuses = {lift: {} for lift in lifts}

# Boshlang'ich komandalar uchun handler
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    markup = ReplyKeyboardMarkup(one_time_keyboard=True)
    for shift in shifts:
        markup.add(KeyboardButton(shift))
    bot.send_message(message.chat.id, "Navbatchilikni tanlang:", reply_markup=markup)

# Navbatchi tanlaganidan keyin lift holatini kiritish uchun
@bot.message_handler(func=lambda message: message.text in shifts)
def choose_lift(message):
    shift = message.text
    bot.send_message(message.chat.id, f"{shift} tanlandi.")
    markup = ReplyKeyboardMarkup(one_time_keyboard=True)
    for lift in lifts:
        markup.add(KeyboardButton(lift))
    bot.send_message(message.chat.id, "Liftni tanlang:", reply_markup=markup)
    bot.register_next_step_handler(message, lambda m: get_lift_status(m, shift))

# Lift holatini kiritish
def get_lift_status(message, shift):
    lift = message.text
    markup = ReplyKeyboardMarkup(one_time_keyboard=True)
    markup.add(KeyboardButton("Ishlayapti"), KeyboardButton("Ishlamayapti"))
    bot.send_message(message.chat.id, f"{lift} uchun holatni tanlang:", reply_markup=markup)
    bot.register_next_step_handler(message, lambda m: save_lift_status(m, shift, lift))

# Kiritilgan ma'lumotlarni saqlash
def save_lift_status(message, shift, lift):
    status = message.text
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lift_statuses[lift][shift] = {'status': status, 'time': current_time}
    bot.send_message(message.chat.id, f"{lift} uchun holat saqlandi: {status}")

# Hujjatni yaratish va yuborish komandasi
@bot.message_handler(commands=['generate_report'])
def generate_report(message):
    doc = Document()
    doc.add_heading('Lift Monitoring Report', 0)

    for lift, statuses in lift_statuses.items():
        doc.add_heading(lift, level=1)
        for shift, data in statuses.items():
            lift_info = f"{shift}: {data['status']} ({data['time']})\n"
            doc.add_paragraph(lift_info)

    filename = "lift_hisobot.docx"
    doc.save(filename)
    with open(filename, 'rb') as f:
        bot.send_document(message.chat.id, f)

# Botni ishga tushirish
bot.polling()
