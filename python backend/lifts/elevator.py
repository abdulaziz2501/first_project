from datetime import datetime
from docx import Document

class LiftMonitoring:
    def __init__(self):
        self.lifts_status = []

    def add_lift_status(self, lift_number, status):
        # Har bir lift holatini vaqt bilan saqlab boradi
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.lifts_status.append({
            'lift_number': lift_number,
            'status': status,
            'time': current_time
        })

    def create_report(self, filename="lift_report.docx"):
        # Hujjat yaratish
        doc = Document()
        doc.add_heading('Lift Monitoring Report', 0)

        for entry in self.lifts_status:
            lift_info = f"Lift: {entry['lift_number']}\nHolati: {entry['status']}\nVaqt: {entry['time']}\n"
            doc.add_paragraph(lift_info)

        doc.save(filename)
        print(f"Hujjat saqlandi: {filename}")

# Dasturdan foydalanish
monitoring = LiftMonitoring()

# Lift holatini qo'shish
monitoring.add_lift_status(1, 'Yaxshi ishlayapti')
monitoring.add_lift_status(2, 'Texnik xizmat kerak')

# Hujjat yaratish
monitoring.create_report("lift_hisobot.docx")
