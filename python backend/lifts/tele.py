import telebot
from datetime import datetime
from docx import Document

API_TOKEN = '7158427390:AAGmiMKa3Fx2TW01-ANC_KHTPSZ9VTc3KDA'  # Bu yerga BotFather dan olingan tokeningizni kiriting

bot = telebot.TeleBot(API_TOKEN)

# Lift holatini saqlash uchun class
class LiftMonitoring:
    def __init__(self):
        self.lifts_status = []

    def add_lift_status(self, lift_number, status):
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.lifts_status.append({
            'lift_number': lift_number,
            'status': status,
            'time': current_time
        })

    def create_report(self, filename="lift_report.docx"):
        doc = Document()
        doc.add_heading('Lift Monitoring Report', 0)

        for entry in self.lifts_status:
            lift_info = f"Lift: {entry['lift_number']}\nHolati: {entry['status']}\nVaqt: {entry['time']}\n"
            doc.add_paragraph(lift_info)

        doc.save(filename)
        return filename

monitoring = LiftMonitoring()

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    bot.reply_to(message, "Salom! Lift monitoring botiga xush kelibsiz! Liftning holatini kiritish uchun /lift_buyrug'ini foydalaning.")

@bot.message_handler(commands=['lift'])
def ask_lift_number(message):
    msg = bot.reply_to(message, "Lift raqamini kiriting:")
    bot.register_next_step_handler(msg, process_lift_number)

def process_lift_number(message):
    try:
        lift_number = int(message.text)
        msg = bot.reply_to(message, f"Lift {lift_number} ning holatini kiriting (masalan, 'Ishlayapti' yoki 'Ishlamayapti'):")
        bot.register_next_step_handler(msg, process_lift_status, lift_number)
    except ValueError:
        bot.reply_to(message, "Iltimos, to'g'ri lift raqamini kiriting.")
        
def process_lift_status(message, lift_number):
    status = message.text
    monitoring.add_lift_status(lift_number, status)
    bot.reply_to(message, f"Lift {lift_number} holati '{status}' deb saqlandi.")

@bot.message_handler(commands=['report'])
def send_report(message):
    report_filename = monitoring.create_report()
    with open(report_filename, 'rb') as doc:
        bot.send_document(message.chat.id, doc)

bot.polling()
