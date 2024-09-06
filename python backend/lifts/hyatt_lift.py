import telebot
from datetime import datetime
from docx import Document

API_TOKEN = '7158427390:AAGmiMKa3Fx2TW01-ANC_KHTPSZ9VTc3KDA'

bot = telebot.TeleBot(API_TOKEN)

# Liftlar nomlari ro'yxati
liftlar = [
    "PL 1", "PL 2", "PL 3", "PL 4", "PL 5",
    "PL 6", "PL 7", "SL 1", "SL 2", "SL 3",
    "SL 4", "SL 5", "SL 6", "SL 7", "SL 8"
]

# Navbatchilar ro'yxati
navbatchilar = {
    '1': "Abdulaziz",
    '2': "Kesha",
    '3': "Ruslan"
}

# Navbatchilar uchun ma'lumotlarni saqlash
navbatchi_data = {key: [] for key in navbatchilar.keys()}

# Ma'lumotlarni qabul qilish funksiyasi
@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "Salom! Qaysi navbatchisiz?\n1 - Abdulaziz\n2- Kesha \n3 - Ruslan \n Navbatchilik raqamingizni kiriting.\n")

@bot.message_handler(func=lambda message: message.text in navbatchilar.keys())
def select_navbatchi(message):
    navbatchi = message.text
    bot.send_message(message.chat.id, f"{navbatchilar[navbatchi]}, liftlar holatini kiriting. Masalan, 'PL 1 ishlayapti' yoki 'SL 1 ishlamayapti'. Barcha liftlar uchun ma'lumot kiriting.")

    @bot.message_handler(func=lambda message: any(lift in message.text for lift in liftlar))
    def get_lift_status(message):
        for lift in liftlar:
            if lift in message.text:
                status = 'Ishlayapti' if 'ishlayapti' in message.text.lower() else 'Ishlamayapti'
                current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                navbatchi_data[navbatchi].append({
                    'lift': lift,
                    'status': status,
                    'time': current_time
                })
                bot.reply_to(message, f"{lift} holati {status} sifatida saqlandi.")

    # Hujjat yaratish va saqlash
    @bot.message_handler(commands=['hujjat'])
    def hujjat(message):
        doc = Document()
        doc.add_heading(f"{navbatchilar[navbatchi]} uchun HYATT Lift Monitoring         Hisoboti", 0)

        for entry in navbatchi_data[navbatchi]:
            lift_info = f"Lift: {entry['lift']}\nHolati: {entry['status']}\nVaqt: {entry['time']}\n"
            doc.add_paragraph(lift_info)

        filename = f"{navbatchilar[navbatchi]}_lift_hisobot_{datetime.now().strftime('%Y_%m_%d')}.docx"
        doc.save(filename)
        bot.send_message(message.chat.id, f"Hujjat saqlandi: {filename}")
        with open(filename, 'rb') as f:
            bot.send_document(message.chat.id, f)

# Botni ishga tushirish
bot.polling()