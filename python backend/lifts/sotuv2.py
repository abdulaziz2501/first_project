import telebot
from telebot import types
from docx import Document

# Bot tokenini o'zingizning tokeningiz bilan almashtiring
TOKEN = '5913724823:AAGLUMyAE1Nb4_UO0G9JnVR77lqagLhn0zY'
bot = telebot.TeleBot(TOKEN)

# Global o'zgaruvchilar
customer_data = {}

# Xaridorning ma'lumotlarini saqlash uchun funksiya
def save_to_word(data):
    document = Document()
    document.add_heading('Xaridor ma\'lumotlari', 0)

    for key, value in data.items():
        document.add_paragraph(f'{key}: {value}')

    document.save(f"{data['Ismi']}_xarid_ma\'lumotlari.docx")

# /start buyrug'i bilan ishga tushadi
@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "Xush kelibsiz! Xaridorning ismini kiriting:")

# Foydalanuvchidan ismni qabul qilish
@bot.message_handler(func=lambda message: True, content_types=['text'])
def get_name(message):
    if 'Ismi' not in customer_data:
        customer_data['Ismi'] = message.text
        bot.reply_to(message, "Telefon raqamini kiriting:")
    elif 'Telefon' not in customer_data:
        customer_data['Telefon'] = message.text
        bot.reply_to(message, "Xarid qilgan mahsulotni kiriting:")
    elif 'Mahsulot' not in customer_data:
        customer_data['Mahsulot'] = message.text
        bot.reply_to(message, "Narxini kiriting:")
    elif 'Narxi' not in customer_data:
        customer_data['Narxi'] = message.text
        save_to_word(customer_data)
        bot.reply_to(message, "Ma'lumotlar saqlandi. Rahmat!")
        customer_data.clear()

# Botni doimiy ravishda ishlatish
bot.polling()
