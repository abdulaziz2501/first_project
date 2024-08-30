import telebot
from telebot.types import ReplyKeyboardMarkup, KeyboardButton
from docx import Document

API_TOKEN = '5913724823:AAGLUMyAE1Nb4_UO0G9JnVR77lqagLhn0zY'  # Bu yerga o'zingizning Telegram bot tokeningizni kiriting

bot = telebot.TeleBot(API_TOKEN)

# Mahsulotlar ro'yxati
products = {
    "Mahsulot 1": {"price": 10000, "description": "Bu mahsulot 1."},
    "Mahsulot 2": {"price": 20000, "description": "Bu mahsulot 2."},
    "Mahsulot 3": {"price": 30000, "description": "Bu mahsulot 3."}
}

# Buyurtmalar ro'yxati
orders = []

# Boshlang'ich komandalar uchun handler
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    markup = ReplyKeyboardMarkup(one_time_keyboard=True)
    markup.add(KeyboardButton("Mahsulotlarni ko'rish"))
    markup.add(KeyboardButton("Buyurtma berish"))
    markup.add(KeyboardButton("Buyurtmalarni ko'rish"))
    bot.send_message(message.chat.id, "Assalomu alaykum! Qanday yordam bera olaman?", reply_markup=markup)

# Mahsulotlarni ko'rsatish
@bot.message_handler(func=lambda message: message.text == "Mahsulotlarni ko'rish")
def show_products(message):
    response = ""
    for product, info in products.items():
        response += f"{product}\nNarxi: {info['price']} so'm\n{info['description']}\n\n"
    bot.send_message(message.chat.id, response)

# Buyurtma berish jarayoni
@bot.message_handler(func=lambda message: message.text == "Buyurtma berish")
def order_product(message):
    markup = ReplyKeyboardMarkup(one_time_keyboard=True)
    for product in products.keys():
        markup.add(KeyboardButton(product))
    bot.send_message(message.chat.id, "Qaysi mahsulotni buyurtma bermoqchisiz?", reply_markup=markup)
    bot.register_next_step_handler(message, confirm_order)

def confirm_order(message):
    product = message.text
    if product in products:
        orders.append({"product": product, "price": products[product]['price'], "buyer": message.chat.first_name})
        bot.send_message(message.chat.id, f"Siz {product} ni buyurtma qildingiz. Buyurtma tasdiqlandi.")
    else:
        bot.send_message(message.chat.id, "Mahsulot topilmadi. Iltimos, qaytadan urinib ko'ring.")

# Buyurtmalarni ko'rsatish va hujjat yaratish
@bot.message_handler(func=lambda message: message.text == "Buyurtmalarni ko'rish")
def show_orders(message):
    if not orders:
        bot.send_message(message.chat.id, "Hozircha buyurtmalar yo'q.")
        return

    doc = Document()
    doc.add_heading('Buyurtmalar ro\'yxati', 0)

    for order in orders:
        order_info = f"Mahsulot: {order['product']}\nNarxi: {order['price']} so'm\nXaridor: {order['buyer']}\n"
        doc.add_paragraph(order_info)

    filename = "buyurtmalar_hisobot.docx"
    doc.save(filename)
    with open(filename, 'rb') as f:
        bot.send_document(message.chat.id, f)

# Botni ishga tushirish
bot.polling()
